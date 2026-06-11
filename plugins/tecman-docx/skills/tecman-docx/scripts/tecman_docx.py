# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Tecman brand palette (teal ramp) ----
TITLE_C = RGBColor(0x00,0x5F,0x66)   # 005F66 title / H2
SUB_C   = RGBColor(0x00,0x7A,0x83)   # 007A83 subtitle / H3
H1_C    = RGBColor(0x1A,0x32,0x34)   # 1A3234 H1 / dark teal
H2_C    = RGBColor(0x00,0x5F,0x66)
H3_C    = RGBColor(0x00,0x7A,0x83)
MUTED   = RGBColor(0x5A,0x7C,0x80)   # 5A7C80 tagline / footer
HDRFILL = "005F66"                   # table header fill
ROWTINT = "E7F1F2"                   # alt row tint
CODEBG  = "F1F5F6"
CODEBORDER = "007A83"
WHITE   = RGBColor(0xFF,0xFF,0xFF)

CW = 9026                            # A4, 1" margins (twips)
import os as _os
LOGO = _os.path.join(_os.path.dirname(_os.path.abspath(__file__)), "..", "assets", "tecman_logo.png")
AUTHOR = "[Author]"
REVIEWERS = "[Reviewers]"
from datetime import date as _date
_t = _date.today()
DATESTR = "%d %s" % (_t.day, _t.strftime("%B %Y"))

_S = {}

def _add_page_field(p, prefix):
    r=p.add_run(prefix); r.font.size=Pt(8); r.font.color.rgb=MUTED; r.font.name='Arial'
    rr=p.add_run(); rr.font.size=Pt(8); rr.font.color.rgb=MUTED; rr.font.name='Arial'
    b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); rr._r.append(b)
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text='PAGE'; rr._r.append(it)
    s=OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'),'separate'); rr._r.append(s)
    e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); rr._r.append(e)

def new_document(doc_title):
    doc=Document()
    n=doc.styles['Normal']; n.font.name='Arial'; n.font.size=Pt(11); n.font.color.rgb=RGBColor(0x20,0x20,0x20)
    sec=doc.sections[0]
    sec.page_width=Mm(210); sec.page_height=Mm(297)
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=Pt(28); sec.footer_distance=Pt(28)
    # running header: small logo, right-aligned
    hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    hp.add_run().add_picture(LOGO, width=Inches(1.25))
    # footer: doc title (left) + page number (right), muted, 8pt
    fp=sec.footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Twips(CW), WD_TAB_ALIGNMENT.RIGHT)
    r=fp.add_run(doc_title + "\t"); r.font.size=Pt(8); r.font.color.rgb=MUTED; r.font.name='Arial'
    _add_page_field(fp, "Page ")
    _S['doc']=doc
    return doc

def _set_shading(el,fill):
    shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill); el.append(shd)
def _cell_shade(cell,fill): _set_shading(cell._tc.get_or_add_tcPr(),fill)

def heading(text,level):
    doc=_S['doc']; p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(14 if level==1 else 10)
    p.paragraph_format.space_after=Pt(6 if level==1 else 4)
    p.paragraph_format.keep_with_next=True
    pPr=p._p.get_or_add_pPr(); ol=OxmlElement('w:outlineLvl'); ol.set(qn('w:val'),str(level-1)); pPr.append(ol)
    r=p.add_run(text); r.bold=True; r.font.name='Arial'
    r.font.size=Pt(16 if level==1 else 13 if level==2 else 12)
    r.font.color.rgb=H1_C if level==1 else H2_C if level==2 else H3_C
    return p

def para(text=None,runs=None,bold=False,italic=False,after=6,before=0,align=None,size=11,color=None):
    doc=_S['doc']; p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    if align: p.alignment=align
    if runs:
        for t,b in runs:
            r=p.add_run(t); r.bold=b; r.font.size=Pt(size); r.font.name='Arial'
            if color is not None: r.font.color.rgb=color
    else:
        r=p.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name='Arial'
        if color is not None: r.font.color.rgb=color
    return p

def bullets(items):
    doc=_S['doc']
    for it in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
        if isinstance(it,list):
            for t,b in it:
                r=p.add_run(t); r.bold=b; r.font.size=Pt(11); r.font.name='Arial'
        else:
            r=p.add_run(it); r.font.size=Pt(11); r.font.name='Arial'

def numbered(items):
    doc=_S['doc']
    for it in items:
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3)
        r=p.add_run(it); r.font.size=Pt(11); r.font.name='Arial'

def code(src):
    doc=_S['doc']; tbl=doc.add_table(rows=1,cols=1); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=tbl.rows[0].cells[0]; _cell_shade(cell,CODEBG)
    cell.width=Twips(CW)
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in src.split('\n'):
        p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0); p.paragraph_format.line_spacing=1.0
        r=p.add_run(line if line else ' '); r.font.name='Consolas'; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x1A,0x1A,0x1A)
        rpr=r._element.get_or_add_rPr(); rf=OxmlElement('w:rFonts'); rf.set(qn('w:ascii'),'Consolas'); rf.set(qn('w:hAnsi'),'Consolas'); rpr.append(rf)
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge,sz in (('left','24'),('top','6'),('bottom','6'),('right','6')):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz); e.set(qn('w:color'),CODEBORDER); b.append(e)
    tcPr.append(b); doc.add_paragraph().paragraph_format.space_after=Pt(4)

def table(headers,rows,widths):
    doc=_S['doc']
    sc=CW/float(sum(widths)); widths=[int(w*sc) for w in widths]
    widths[-1]=CW-sum(widths[:-1])
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.allow_autofit=False
    hdr=t.rows[0].cells
    for i,h in enumerate(headers):
        _cell_shade(hdr[i],HDRFILL); p=hdr[i].paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(2)
        r=p.add_run(h); r.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(10); r.font.name='Arial'
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            _cell_shade(cells[i],'FFFFFF' if ri%2 else ROWTINT); p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(2)
            r=p.add_run(val); r.font.size=Pt(10); r.font.name='Arial'
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Twips(w)
    tblPr=t._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'C7D7D9'); b.append(e)
    tblPr.append(b); doc.add_paragraph().paragraph_format.space_after=Pt(4)

def page_break():
    _S['doc'].add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def control_block(title, code_, subtitle, version, status, scope, related, change):
    doc=_S['doc']
    # large logo, left
    lp=doc.add_paragraph(); lp.paragraph_format.space_after=Pt(10); lp.paragraph_format.space_before=Pt(6)
    lp.add_run().add_picture(LOGO, width=Inches(3.0))
    p=para(title,bold=True,before=10,after=2,size=24,color=TITLE_C)
    p=para(subtitle,after=2,size=12,color=SUB_C)
    para("Tecman — AI Handy Mapper",italic=True,after=14,size=11,color=MUTED)
    table(['Field','Detail'],[
        ['Document',title],
        ['Document code',code_],
        ['Version',version],
        ['Status',status],
        ['Date',DATESTR],
        ['Author',AUTHOR],
        ['Reviewers',REVIEWERS],
        ['Approver','[TBC]'],
        ['Scope / phase',scope],
        ['Related documents',related],
    ],[2400,6626])
    heading('Change History',2)
    table(['Version','Date','Author','Summary'], change, [1150,1500,2700,3676])
    page_break()
    heading('Contents',1)
    p=doc.add_paragraph(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'TOC \\o "1-2" \\h \\z \\u')
    it=OxmlElement('w:r'); itt=OxmlElement('w:t'); itt.text='Right-click and choose "Update Field" to build the contents.'; it.append(itt); fld.append(it); p._p.append(fld)
    page_break()

def save(path):
    _S['doc'].save(path); print('saved', path.split('/')[-1])

def figure(path, width_in, caption):
    doc=_S['doc']
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    p.add_run().add_picture(path, width=Inches(width_in))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(8)
    r=c.add_run(caption); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=MUTED; r.font.name='Arial'
